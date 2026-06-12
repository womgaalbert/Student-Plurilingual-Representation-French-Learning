"""
generate_targets_explanation.py
Explication complète des cibles H1-H2-H3-H4 pour la chercheuse
Usage : python scripts/generate_targets_explanation.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path("reports/explication_des_cibles_FLP.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)


# ── helpers ───────────────────────────────────────────────────────────────────

def cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    colors = {1: "1F497D", 2: "2E74B5", 3: "404040"}
    sizes  = {1: 15, 2: 13, 3: 11}
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(colors.get(level, "000000"))
    run.font.size = Pt(sizes.get(level, 11))
    return p


def table(doc, rows, header, col_widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        cell_bg(hdr[i], "1F497D")
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        if col_widths:
            hdr[i].width = Inches(col_widths[i])
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            p = row[i].paragraphs[0]
            if p.runs:
                p.runs[0].font.size = Pt(9)
            if col_widths:
                row[i].width = Inches(col_widths[i])
    return t


def code(doc, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Code"] if "Code" in [s.name for s in doc.styles] else doc.styles["Normal"]
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    return p


def bold_para(doc, label: str, text: str):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)
    return p


def footer(doc):
    sec = doc.sections[0]
    ft = sec.footer
    p = ft.paragraphs[0]
    p.text = "FLP — Explication des cibles ML par hypothèse — Juin 2026"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── document ──────────────────────────────────────────────────────────────────

doc = Document()
sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.5)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)
footer(doc)

# ── page de garde ─────────────────────────────────────────────────────────────

for _ in range(3): doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("EXPLICATION DES CIBLES ML\nPAR HYPOTHÈSE")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Projet FLP — French Learning Perceptions\nin Plurilingual Cameroon")
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
for label, val in [
    ("Chercheuse", "Chancelline Armelle Nongni Kendjio\nDoctorante en Didactique du FLES\nUniversité Marie & Louis Pasteur de Besançon"),
    ("Support ML/AI", "Albert Womga"),
    ("Date", "Juin 2026"),
    ("Document", "Explication pédagogique des cibles ML"),
]:
    p = doc.add_paragraph()
    p.add_run(f"{label} : ").bold = True
    p.add_run(val)

doc.add_page_break()

# ── introduction ──────────────────────────────────────────────────────────────

heading(doc, "Introduction", 1)
doc.add_paragraph(
    "Ce document explique, dans un langage accessible à une chercheuse en didactique du FLES, "
    "comment chaque hypothèse de la thèse est traduite en une ou plusieurs cibles (variables à prédire) "
    "pour les modèles de Machine Learning."
)
doc.add_paragraph(
    "Pour chaque hypothèse, vous trouverez : la ou les questions du questionnaire qui servent de source, "
    "la règle de calcul de la cible (en langage naturel ET en pseudo-code), et ce que chaque valeur "
    "de la cible signifie concrètement sur le terrain — c'est-à-dire pour vos élèves."
)
doc.add_paragraph(
    "Les quatre hypothèses sont présentées dans l'ordre du pipeline : H1, H2, H3, H4."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# H1
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Hypothèse H1 — Répertoire Multilingue & Mobilisation des Langues", 1)

doc.add_paragraph(
    "Hypothèse : Le contexte multilingue dans lequel évoluent les élèves, ainsi que la diversité "
    "de leur répertoire linguistique, les amènent à mobiliser plusieurs langues dans l'ensemble "
    "de leurs interactions quotidiennes."
)

heading(doc, "Question source", 2)
bold_para(doc, "Question du questionnaire : ",
          "« Utilisez-vous d'autres langues dans votre quotidien ? »")
bold_para(doc, "Colonne technique : ", "usage_quotidien")
bold_para(doc, "Variable cible : ", "h1_target")

heading(doc, "Règle de calcul", 2)
doc.add_paragraph(
    "La réponse de l'élève est analysée pour détecter si elle commence par OUI ou par NON."
)
table(doc, [
    ("OUI = 1", "La réponse commence par « Oui »"),
    ("NON = 0", "La réponse commence par « Non »"),
    ("-1 (exclu)", "La réponse est absente ou ambiguë — l'élève est retiré de l'analyse"),
], ("Valeur", "Règle"))

heading(doc, "Ce que la cible signifie concrètement", 2)

bold_para(doc, "h1_target = 1 (OUI) : ",
          "L'élève déclare utiliser activement d'autres langues que le français dans sa vie "
          "de tous les jours. Cela signifie qu'il mobilise réellement son répertoire plurilingue — "
          "il parle ewondo à la maison, foufouldé au marché, anglais avec des amis, etc. "
          "Son plurilinguisme n'est pas seulement théorique : il est mis en pratique au quotidien.")

bold_para(doc, "h1_target = 0 (NON) : ",
          "L'élève déclare ne pas utiliser d'autres langues au quotidien. Cela ne signifie pas "
          "nécessairement qu'il ne connaît pas d'autres langues — il peut avoir un répertoire "
          "plurilingue mais ne pas le mobiliser dans ses interactions journalières, par exemple "
          "parce que son environnement immédiat est majoritairement francophone.")

heading(doc, "Tâche ML associée", 2)
table(doc, [
    ("Type", "Classification binaire"),
    ("Modèle", "XGBoost + SMOTE + StandardScaler"),
    ("Ce que le modèle apprend", "À prédire si un élève mobilise ses langues au quotidien (OUI/NON) "
     "à partir de son profil linguistique (nombre de langues, fréquence d'exposition, etc.)"),
], ("Aspect", "Description"))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# H2
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Hypothèse H2 — Représentations du Français → Motivation & Difficultés", 1)

doc.add_paragraph(
    "Hypothèse : Les représentations du français influencent la motivation et les difficultés "
    "d'apprentissage des élèves."
)
doc.add_paragraph(
    "H2 est la seule hypothèse à avoir DEUX cibles distinctes, car elle cherche à capturer "
    "deux facettes de la relation élève-français : la motivation (A) et les obstacles perçus (B)."
)

# ── H2 - Cible A ─────────────────────────────────────────────────────────────

heading(doc, "Cible A — Motivation à apprendre le français", 2)

heading(doc, "Question source", 3)
bold_para(doc, "Question du questionnaire : ",
          "« Qu'est-ce qui vous motive à apprendre le français ? »")
bold_para(doc, "Colonne technique : ", "motivation_apprendre")
bold_para(doc, "Variable cible : ", "h2_target_motivation")

heading(doc, "Règle de calcul", 3)
doc.add_paragraph(
    "La réponse (texte libre) est analysée pour détecter la présence de mots-clés "
    "qui révèlent une motivation élevée ou au contraire une motivation faible."
)

table(doc, [
    ("Motivation ÉLEVÉE", "communiquer, réussir, avenir, s'exprimer, aider, voyager, études, comprendre"),
    ("Motivation FAIBLE", "obligé, forcé, difficile, rien"),
], ("Catégorie", "Mots-clés recherchés"))

doc.add_paragraph()
doc.add_paragraph("La règle de décision est la suivante :")

table(doc, [
    ("2", "Motivation ÉLEVÉE", "La réponse contient ≥ 2 mots-clés de motivation élevée",
     "« Communiquer avec les étrangers, réussir mes études et voyager » → 3 mots-clés"),
    ("1", "Motivation MOYENNE", "Aucun des deux seuils n'est atteint (ou réponse absente)",
     "« Parce que c'est la langue officielle » → 0 mot-clé dans les deux listes"),
    ("0", "Motivation FAIBLE", "La réponse contient ≥ 1 mot-clé de motivation faible",
     "« Je suis obligé, c'est difficile » → obligé + difficile = 2 mots-clés"),
], ("Valeur", "Label", "Condition", "Exemple de réponse d'élève"))

heading(doc, "Ce que la cible signifie concrètement", 3)

bold_para(doc, "h2_target_motivation = 2 : ",
          "L'élève exprime une motivation intrinsèque forte. Il perçoit le français comme "
          "un outil d'avenir — pour communiquer, étudier, voyager, réussir professionnellement. "
          "Il n'apprend pas le français par obligation, mais par désir personnel.")

bold_para(doc, "h2_target_motivation = 1 : ",
          "L'élève est dans une position intermédiaire. Sa motivation n'est ni particulièrement "
          "forte ni particulièrement faible. C'est le profil le plus fréquent : l'élève accepte "
          "l'apprentissage du français sans enthousiasme marqué ni rejet explicite.")

bold_para(doc, "h2_target_motivation = 0 : ",
          "L'élève est en motivation extrinsèque contrainte. Il apprend le français parce qu'il "
          "y est obligé, non par choix. Cette posture est souvent associée à un discours de "
          "difficulté ou de résignation.")

# ── H2 - Cible B ─────────────────────────────────────────────────────────────

heading(doc, "Cible B — Types de difficultés perçues (multi-label)", 2)

heading(doc, "Questions sources", 3)
doc.add_paragraph(
    "Contrairement à la cible A qui repose sur une seule question, la cible B fusionne "
    "TROIS questions du questionnaire pour capturer tous les aspects des difficultés "
    "que l'élève rencontre :"
)
table(doc, [
    ("1", "« Quels aspects du français trouvez-vous faciles ou difficiles à apprendre ? [Difficile] »"),
    ("2", "« Quelles sont les principales difficultés que vous rencontrez en apprenant le français ? »"),
    ("3", "« Est-ce que ces difficultés sont liées à la langue elle-même, au vocabulaire, à la grammaire ou à d'autres aspects ? »"),
], ("N°", "Question"))

bold_para(doc, "Variable cible : ", "diff_{domaine} (7 variables binaires)")

heading(doc, "Règle de calcul", 3)
doc.add_paragraph(
    "Les trois réponses sont concaténées en un seul texte. Pour chaque domaine de difficulté, "
    "on détecte si des mots-clés spécifiques apparaissent dans ce texte combiné."
)

table(doc, [
    ("diff_grammaire", "La réponse mentionne la grammaire",
     "grammaire, grammatical, accord"),
    ("diff_vocabulaire", "La réponse mentionne le vocabulaire/lexique",
     "vocabulaire, mots, lexique"),
    ("diff_orthographe", "La réponse mentionne l'orthographe",
     "orthographe, fautes"),
    ("diff_conjugaison", "La réponse mentionne la conjugaison",
     "conjugaison, conjuguer, verbes"),
    ("diff_expression_orale", "La réponse mentionne l'oral",
     "expression orale, oral, parler"),
    ("diff_comprehension", "La réponse mentionne la compréhension/lecture",
     "compréhension, comprendre, lecture"),
    ("diff_analyse", "La réponse mentionne l'analyse",
     "analyse, analyser"),
], ("Variable", "Ce que 1 signifie", "Mots-clés détectés"))

heading(doc, "Ce que la cible signifie concrètement", 3)
doc.add_paragraph(
    "Un même élève peut avoir plusieurs diff_{domaine} = 1. Par exemple, un élève qui écrit : "
    "« J'ai du mal avec la grammaire et la conjugaison des verbes » aura :"
)
doc.add_paragraph("• diff_grammaire = 1", style="List Bullet")
doc.add_paragraph("• diff_conjugaison = 1", style="List Bullet")
doc.add_paragraph("• diff_vocabulaire = 0, diff_orthographe = 0, diff_expression_orale = 0, diff_comprehension = 0, diff_analyse = 0", style="List Bullet")

doc.add_paragraph(
    "Le modèle ClassifierChain utilisé pour cette cible capture le fait que certaines difficultés "
    "vont souvent ensemble — par exemple, un élève qui mentionne la grammaire mentionne aussi "
    "souvent la conjugaison. Ces dépendances entre labels sont exploitées par le modèle."
)

heading(doc, "Tâche ML associée (H2 global)", 2)
table(doc, [
    ("Cible A — Type", "Classification 3 classes (0/1/2)"),
    ("Cible A — Modèle", "XGBoost + RandomOverSampler"),
    ("Cible B — Type", "Classification multi-label (7 labels binaires)"),
    ("Cible B — Modèle", "ClassifierChain(XGBoost) — capture les dépendances entre difficultés"),
    ("Ce que les modèles apprennent",
     "A : À prédire le niveau de motivation à partir de la perception du français.\n"
     "B : À identifier les types de difficultés à partir des déclarations de l'élève."),
], ("Aspect", "Description"))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# H3
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Hypothèse H3 — Exposition Plurilingue → Attitudes envers le Français", 1)

doc.add_paragraph(
    "Hypothèse : L'exposition au plurilinguisme influence positivement les attitudes "
    "envers le français."
)
doc.add_paragraph(
    "H3 se distingue des autres hypothèses par son approche : elle combine une tâche de "
    "régression, une tâche de classification, ET un test d'inférence causale. "
    "La cible est un SCORE COMPOSITE construit à partir de deux questions."
)

heading(doc, "Questions sources", 2)
table(doc, [
    ("Composante s1", "« Comment le français se compare-t-il aux autres langues que vous connaissez ou que vous entendez dans votre environnement ? »",
     "Perception comparative du français"),
    ("Composante s2", "« Pensez-vous que l'apprentissage du français en même temps que les langues camerounaises vous motiverait ? »",
     "Motivation plurilingue scolaire"),
], ("Nom technique", "Question du questionnaire", "Ce qu'elle mesure"))

heading(doc, "Construction du score composite — Règle de calcul", 2)
doc.add_paragraph(
    "Le score est un nombre entre 1,0 et 5,0 qui fusionne les deux composantes avec "
    "des poids différents."
)

heading(doc, "Étape 1 — Calcul de s1 (poids 60%)", 3)
table(doc, [
    ("3.0", "L'élève dit que le français est PLUS important que les autres langues",
     "« Le français est plus important parce que c'est une langue internationale »"),
    ("2.0", "L'élève dit que le français est AUTANT important que les autres langues",
     "« C'est aussi important, chaque langue a sa valeur »"),
    ("1.0", "L'élève dit que le français est MOINS important que les autres langues",
     "« Ma langue maternelle est plus importante pour moi »"),
], ("Score s1", "Signification", "Exemple"))

heading(doc, "Étape 2 — Calcul de s2 (poids 40%)", 3)
table(doc, [
    ("2.0", "Très bien — L'élève serait très motivé"),
    ("1.5", "Bien — L'élève serait plutôt motivé"),
    ("1.0", "Un peu — L'élève serait peu motivé"),
    ("0.5", "Pas du tout — L'élève ne serait pas motivé"),
], ("Score s2", "Réponse de l'élève"))

heading(doc, "Étape 3 — Score final h3_score_attitude", 3)
doc.add_paragraph(
    "Le score combine les deux composantes selon une moyenne pondérée à 60/40, "
    "puis normalise le résultat sur l'échelle [1,0 – 5,0]. "
    "Plus le score est élevé, plus l'élève a une attitude positive envers le français "
    "dans son contexte plurilingue."
)

heading(doc, "Classification associée — h3_attitude_class", 2)
table(doc, [
    ("Positive", "≥ 3,5",
     "L'élève valorise le français dans son écosystème plurilingue. Il ne voit pas le français "
     "et les langues locales comme antagonistes, mais comme complémentaires."),
    ("Neutre", "2,5 – 3,4",
     "Position ambivalente — l'élève n'exprime ni enthousiasme marqué ni rejet."),
    ("Négative", "< 2,5",
     "L'élève minimise l'importance relative du français ou rejette l'idée d'un apprentissage "
     "plurilingue."),
], ("Classe", "Score brut", "Ce que cela révèle sur l'élève"))

heading(doc, "Volet causal (spécifique à H3)", 2)
doc.add_paragraph(
    "Au-delà du ML, H3 comporte une analyse causale qui teste si l'exposition aux autres "
    "langues (exposition_bin) est corrélée au score d'attitude. L'ATE (Average Treatment Effect) "
    "mesure la différence moyenne de score entre les élèves exposés et non-exposés."
)
doc.add_paragraph(
    "Si l'ATE est proche de zéro et que le test n'est pas significatif, cela ne signifie pas "
    "que l'hypothèse est fausse — cela peut simplement révéler que l'exposition est quasi-universelle "
    "dans l'échantillon (ce qui est une découverte en soi)."
)

heading(doc, "Tâche ML associée", 2)
table(doc, [
    ("Régression", "RandomForest Regressor → Prédit h3_score_attitude sur [1–5]"),
    ("Classification", "XGBoost Classifier → Prédit h3_attitude_class (Positive/Neutre/Négative)"),
    ("Causal", "Pearson r + ATE → Teste exposition_bin → attitude"),
], ("Volet", "Modèle"))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# H4
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Hypothèse H4 — Intégration Langues Locales → Engagement & Motivation", 1)

doc.add_paragraph(
    "Hypothèse : L'utilisation des langues locales pendant les activités d'enseignement "
    "renforce l'intérêt pour les tâches au cours de français et améliore l'engagement "
    "global dans le cours."
)
doc.add_paragraph(
    "H4 est l'hypothèse la plus complexe : elle comporte TROIS cibles distinctes, "
    "chacune capturant une dimension différente du rapport des élèves à l'intégration "
    "des langues camerounaises en classe."
)

# ── H4 - Cible A ─────────────────────────────────────────────────────────────

heading(doc, "Cible A — Motivation pour l'intégration des langues locales (binaire)", 2)

heading(doc, "Question source", 3)
bold_para(doc, "Question du questionnaire : ",
          "« Pensez-vous que l'apprentissage du français en même temps que les langues "
          "camerounaises vous motiverait ? »")
bold_para(doc, "Colonne technique : ", "motivation_camerounaises")
bold_para(doc, "Variable cible : ", "h4_target_motivation")

heading(doc, "Règle de calcul", 3)
doc.add_paragraph(
    "Les réponses sont sur une échelle à 4 niveaux (Très bien / Bien / Un peu / Pas du tout). "
    "On les regroupe en deux catégories."
)

table(doc, [
    ("1 (OUI)", "Très bien", "3",
     "L'élève serait TRÈS motivé par l'intégration des langues camerounaises dans le cours de français"),
    ("1 (OUI)", "Bien", "2",
     "L'élève serait PLUTÔT motivé — il est ouvert à l'idée sans être le plus enthousiaste"),
    ("0 (NON)", "Un peu", "1",
     "L'élève serait PEU motivé — l'intégration ne changerait pas grand-chose pour lui"),
    ("0 (NON)", "Pas du tout", "0",
     "L'élève ne serait PAS DU TOUT motivé — il est opposé ou indifférent à cette idée"),
], ("Cible", "Réponse brute", "Score INTERET_MAP", "Ce que cela signifie"))

heading(doc, "Ce que la cible signifie concrètement", 3)

bold_para(doc, "h4_target_motivation = 1 : ",
          "Cet élève déclare qu'apprendre le français en même temps que les langues "
          "camerounaises le motiverait davantage. C'est un signal fort : l'élève perçoit "
          "l'intégration des langues locales non pas comme une distraction, mais comme "
          "un levier de motivation supplémentaire.")

bold_para(doc, "h4_target_motivation = 0 : ",
          "Cet élève ne pense pas que l'intégration des langues camerounaises changerait "
          "sa motivation. Cela peut refléter une indifférence, une préférence pour un "
          "enseignement 100% en français, ou simplement l'absence d'opinion sur le sujet.")

# ── H4 - Cible B ─────────────────────────────────────────────────────────────

heading(doc, "Cible B — Score d'engagement (ordinal 1–4)", 2)

heading(doc, "Questions sources", 3)
doc.add_paragraph(
    "Ce score composite fusionne DEUX questions pour mesurer non seulement si l'élève "
    "EST motivé, mais aussi à quel FRÉQUENCE il souhaiterait cette intégration."
)

table(doc, [
    ("1", "« Pensez-vous que l'apprentissage du français en même temps que les langues camerounaises vous motiverait ? »",
     "Capte l'INTENSITÉ de la motivation (0 ou 2 points)"),
    ("2", "« Aimeriez-vous que les cours de français incluent davantage d'activités sur d'autres langues ? »",
     "Capte la FRÉQUENCE souhaitée (0 à 2 points, selon échelle Toujours/Souvent/Parfois/Rarement/Jamais)"),
], ("N°", "Question", "Rôle dans le score"))

heading(doc, "Règle de calcul — La formule", 3)

doc.add_paragraph(
    "Le score d'engagement combine l'intensité de la motivation (m) et la fréquence "
    "souhaitée d'intégration (f) en un score unique entre 1 et 4."
)

doc.add_paragraph(
    "Composante m (motivation) : vaut 2 si l'élève répond « Très bien » ou « Bien », 0 sinon."
)
doc.add_paragraph(
    "Composante f (fréquence) : Toujours=4, Souvent=3, Parfois=2, Rarement=1, Jamais=0. "
    "Puis normalisé sur [0, 2] via la formule f_norm = (f / 4) × 2."
)
doc.add_paragraph(
    "Score brut = m + f_norm, plage [0, 4]. Ce brut est ensuite discrétisé :"
)

table(doc, [
    ("4 — Engagement MAXIMAL", "≥ 3,5",
     "L'élève est motivé ET souhaite une intégration très fréquente (Toujours ou Souvent)"),
    ("3 — Engagement FORT", "≥ 2,5",
     "L'élève est motivé avec une fréquence modérée, OU non motivé mais souhaite une intégration fréquente"),
    ("2 — Engagement MODÉRÉ", "≥ 1,0",
     "L'élève est motivé mais ne souhaite pas d'intégration fréquente, OU l'inverse"),
    ("1 — Engagement FAIBLE", "< 1,0",
     "L'élève n'est ni motivé, ni demandeur d'intégration"),
], ("Score", "Seuil brut", "Profil de l'élève"))

heading(doc, "Ce que la cible signifie concrètement", 3)

bold_para(doc, "h4_engagement_score = 4 : ",
          "Profil idéal — l'élève est non seulement favorable à l'intégration des langues "
          "locales, mais il souhaite qu'elle soit fréquente. C'est un « ambassadeur » potentiel "
          "de l'approche plurilingue dans sa classe.")

bold_para(doc, "h4_engagement_score = 1 : ",
          "L'élève n'exprime ni motivation pour l'intégration des langues locales, ni souhait "
          "de fréquence. Il est en retrait par rapport à cette proposition pédagogique.")

# ── H4 - Cible C ─────────────────────────────────────────────────────────────

heading(doc, "Cible C — Disciplines souhaitées pour l'intégration (multi-label)", 2)

heading(doc, "Question source", 3)
bold_para(doc, "Question du questionnaire : ",
          "« À quelle discipline du français aimeriez-vous que l'apprentissage des langues "
          "camerounaises soit associé ? »")
bold_para(doc, "Colonne technique : ", "discipline_associee")
bold_para(doc, "Variable cible : ", "vd_disc_{discipline} (5 variables binaires)")

heading(doc, "Règle de calcul", 3)
doc.add_paragraph(
    "La réponse de l'élève est analysée pour détecter quelle(s) discipline(s) il mentionne. "
    "Un élève peut mentionner plusieurs disciplines."
)

table(doc, [
    ("vd_disc_vocabulaire", "L'élève souhaite les langues locales en VOCABULAIRE",
     "vocabulaire, mots, lexique"),
    ("vd_disc_grammaire", "L'élève souhaite les langues locales en GRAMMAIRE",
     "grammaire"),
    ("vd_disc_lecture", "L'élève souhaite les langues locales en LECTURE",
     "lecture, lire, textes, récit"),
    ("vd_disc_expression_orale", "L'élève souhaite les langues locales en EXPRESSION ORALE",
     "expression orale, oral, parler"),
    ("vd_disc_conjugaison", "L'élève souhaite les langues locales en CONJUGAISON",
     "conjugaison, conjuguer"),
], ("Variable", "Ce que 1 signifie", "Mots-clés détectés"))

heading(doc, "Ce que la cible signifie concrètement", 3)
doc.add_paragraph(
    "Cette cible permet de répondre à la question : dans quelle(s) discipline(s) les élèves "
    "souhaitent-ils prioritairement que les langues camerounaises soient intégrées ?"
)
doc.add_paragraph(
    "Le ClassifierChain classe ensuite ces disciplines par ordre de priorité (via la moyenne "
    "des probabilités prédites), ce qui donne une recommandation pédagogique actionable : "
    "« Commencez par intégrer les langues locales en conjugaison et vocabulaire comparatif, "
    "disciplines pour lesquelles la demande est la plus forte. »"
)

heading(doc, "Tâche ML associée (H4 global)", 2)
table(doc, [
    ("Cible A — Type", "Classification binaire"),
    ("Cible A — Modèle", "XGBoost avec scale_pos_weight"),
    ("Cible B — Type", "Classification ordinale (4 niveaux)"),
    ("Cible B — Modèle", "XGBoost multi-classe (num_class=4)"),
    ("Cible C — Type", "Classification multi-label (5 disciplines)"),
    ("Cible C — Modèle", "ClassifierChain(XGBoost) — capture les combinaisons de disciplines"),
    ("Ce que les modèles apprennent",
     "A : À prédire si un élève serait motivé par l'intégration.\n"
     "B : À prédire le niveau d'engagement composite.\n"
     "C : À identifier les disciplines prioritaires pour l'intégration."),
], ("Aspect", "Description"))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SYNTHÈSE
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Tableau de synthèse — Les 8 cibles du projet FLP", 1)

table(doc, [
    ("H1", "h1_target", "Binaire (0/1)", "Utilisez-vous d'autres langues dans votre quotidien ?",
     "Mobilisation réelle du répertoire plurilingue"),
    ("H2", "h2_target_motivation", "3 classes (0/1/2)", "Qu'est-ce qui vous motive à apprendre le français ?",
     "Niveau de motivation (faible/moyen/élevé)"),
    ("H2", "diff_{domaine} ×7", "Multi-label (0/1)", "3 questions fusionnées sur les difficultés",
     "Types de difficultés identifiés par l'élève"),
    ("H3", "h3_score_attitude", "Régression [1-5]", "Comparaison + Motivation (2 questions)",
     "Score composite d'attitude envers le français en contexte plurilingue"),
    ("H3", "h3_attitude_class", "3 classes", "Dérivé du score ci-dessus",
     "Positive / Neutre / Négative"),
    ("H4", "h4_target_motivation", "Binaire (0/1)", "Seriez-vous motivé par français + langues camerounaises ?",
     "Motivation pour l'intégration des langues locales"),
    ("H4", "h4_engagement_score", "Ordinal [1-4]", "Motivation + Souhait de fréquence (2 questions)",
     "Niveau d'engagement composite"),
    ("H4", "vd_disc_{x} ×5", "Multi-label (0/1)", "Dans quelle discipline intégrer les langues locales ?",
     "Disciplines souhaitées pour l'intégration"),
], ("Hyp.", "Variable cible", "Type ML", "Question(s) source", "Ce qu'elle mesure"))

doc.add_paragraph()
heading(doc, "Note importante pour la chercheuse", 2)
doc.add_paragraph(
    "Plusieurs cibles sont des scores COMPOSITES (H3, H4-B), c'est-à-dire construits par "
    "règles à partir de plusieurs questions du questionnaire. Ces cibles ne sont pas des "
    "mesures directes mais des proxys — elles approximent un construit théorique "
    "(l'attitude, l'engagement) à partir des réponses disponibles."
)
doc.add_paragraph(
    "C'est une limite assumée de cette première phase quantitative. La validité de ces proxys "
    "repose sur la cohérence interne des réponses des élèves et sur les cadres théoriques "
    "mobilisés (Moscovici & Jodelet pour les représentations, échelles de Likert pour les "
    "attitudes). Une phase qualitative complémentaire (entretiens semi-directifs) permettrait "
    "de valider ou d'affiner ces construits."
)

# ── footer ────────────────────────────────────────────────────────────────────

doc.save(str(OUT))
print(f"Document sauvegardé : {OUT}")
print(f"Taille : {OUT.stat().st_size / 1024:.0f} Ko")
