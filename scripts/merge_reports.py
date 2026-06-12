"""
merge_reports.py — Fusionne explication_des_cibles_FLP.docx dans rapport_complet_FLP_mai2026.docx
Usage : python scripts/merge_reports.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy

REPORTS = Path("reports")
TARGET  = REPORTS / "rapport_complet_FLP_mai2026.docx"
SOURCE  = REPORTS / "explication_des_cibles_FLP.docx"
OUT     = REPORTS / "rapport_complet_FLP_mai2026.docx"  # overwrite

# ── Helper: Copy an element fully ─────────────────────────────────────────────

def copy_element(elm):
    """Deep-copy an lxml element, clearing rsid attributes."""
    dup = deepcopy(elm)
    # Clear revision IDs to avoid conflicts
    for attr in list(dup.attrib):
        if attr.endswith("rsidR") or attr.endswith("rsidRPr") or attr.endswith("rsidP") or attr == "rsidRDefault":
            del dup.attrib[attr]
    for desc in dup.iter():
        for attr in list(desc.attrib):
            if attr.endswith("rsidR") or attr.endswith("rsidRPr") or attr.endswith("rsidP") or attr == "rsidRDefault":
                del desc.attrib[attr]
    return dup


def copy_paragraph(src_para, tgt_doc):
    """Copy a paragraph from source to target document, preserving style and runs."""
    new_p = tgt_doc.add_paragraph()
    # Copy paragraph style
    if src_para.style:
        try:
            new_p.style = src_para.style
        except Exception:
            pass
    new_p.alignment = src_para.alignment

    for run in src_para.runs:
        new_run = new_p.add_run(run.text)
        if run.bold:       new_run.bold = run.bold
        if run.italic:     new_run.italic = run.italic
        if run.font.size:  new_run.font.size = run.font.size
        if run.font.name:  new_run.font.name = run.font.name
        if run.font.color and run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
        if run.underline:  new_run.underline = run.underline
    return new_p


def copy_table(src_table, tgt_doc):
    """Copy a table from source to target document."""
    rows_n = len(src_table.rows)
    cols_n = len(src_table.columns)
    new_table = tgt_doc.add_table(rows=rows_n, cols=cols_n)
    new_table.style = src_table.style

    for ri, src_row in enumerate(src_table.rows):
        for ci, src_cell in enumerate(src_row.cells):
            tgt_cell = new_table.rows[ri].cells[ci]
            # Copy cell text
            for pi, para in enumerate(src_cell.paragraphs):
                if pi == 0:
                    tgt_para = tgt_cell.paragraphs[0]
                    # Clear default empty run
                    if tgt_para.runs:
                        tgt_para.clear()
                else:
                    tgt_para = tgt_cell.add_paragraph()
                for run in para.runs:
                    nr = tgt_para.add_run(run.text)
                    if run.bold:        nr.bold = run.bold
                    if run.italic:      nr.italic = run.italic
                    if run.font.size:   nr.font.size = run.font.size
                    if run.font.name:   nr.font.name = run.font.name
                    if run.font.color and run.font.color.rgb:
                        nr.font.color.rgb = run.font.color.rgb
                    if run.underline:   nr.underline = run.underline
            # Copy cell shading
            src_tc = src_cell._tc
            tgt_tc = tgt_cell._tc
            src_shd = src_tc.findall(qn("w:tcPr") + "/" + qn("w:shd"), src_tc.nsmap)
            if src_shd:
                tgt_tcPr = tgt_tc.get_or_add_tcPr()
                for shd in src_shd:
                    tgt_tcPr.append(deepcopy(shd))
    return new_table


# ── Merge ─────────────────────────────────────────────────────────────────────

print("Chargement du document cible...")
tgt = Document(str(TARGET))

print("Chargement du document source...")
src = Document(str(SOURCE))

# Find insertion point: after section 2 "Méthodologie" and before section 3
# We look for the paragraph containing "3. Nettoyage et description des données"
insert_idx = None
for i, para in enumerate(tgt.paragraphs):
    if para.style.name.startswith("Heading") and "3." in para.text:
        insert_idx = i
        break

if insert_idx is None:
    # Fallback: insert after "2. Méthodologie"
    for i, para in enumerate(tgt.paragraphs):
        if para.style.name.startswith("Heading") and "Méthodologie" in para.text:
            insert_idx = i + 1
            # Skip until the next heading
            for j in range(insert_idx, len(tgt.paragraphs)):
                if tgt.paragraphs[j].style.name.startswith("Heading"):
                    insert_idx = j
                    break
            break

if insert_idx is None:
    # Absolute fallback: insert before the last page break / before synthèse
    insert_idx = max(1, len(tgt.paragraphs) - 5)

print(f"Insertion apres le paragraphe {insert_idx}")

# Determine which body element to insert after
# We need to work at the XML body level
body = tgt.element.body
target_para_xml = tgt.paragraphs[insert_idx]._element

from docx.enum.text import WD_BREAK

# Insert a page break before the new section
pb = tgt.add_paragraph()
pb.add_run().add_break(WD_BREAK.PAGE)
pb_el = pb._element

# Find the index of the target paragraph in the body
body_children = list(body)
target_idx = None
for i, child in enumerate(body_children):
    if child is target_para_xml:
        target_idx = i
        break

# Move the page break right after the target paragraph
if target_idx is not None:
    body.remove(pb_el)
    body.insert(target_idx + 1, pb_el)
    insert_after_idx = target_idx + 1  # insert after the page break
else:
    insert_after_idx = len(body) - 1

# Now insert all source body elements after the page break
# We skip the first few elements of the source (cover page, introduction up to "Les quatre hypothèses")
# and insert everything from H1 explanation onwards

skipped = 0
started = False
inserted = 0
for src_child in list(src.element.body):
    tag = src_child.tag.split("}")[-1] if "}" in src_child.tag else src_child.tag

    # Detect when we reach H1 section (Heading 1 with "Hypothèse H1")
    if tag == "p":
        pPr = src_child.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and "Heading" in (pStyle.get(qn("w:val")) or ""):
                texts = []
                for t in src_child.iter(qn("w:t")):
                    texts.append(t.text or "")
                full = "".join(texts)
                if "Hypothèse H1" in full:
                    started = True

    if not started:
        skipped += 1
        continue

    # Copy this element
    dup = copy_element(src_child)
    body.insert(insert_after_idx + 1 + inserted, dup)
    inserted += 1

print(f"Elements sautes (couverture + intro source) : {skipped}")
print(f"Elements inseres (corps H1 -> Synthese) : {inserted}")

# Add a note before the inserted section
note_p = tgt.paragraphs[insert_idx]
# Actually, let's add a section title before the inserted content
# We inserted the page break at insert_after_idx, then content after that
# Let's add a header right after the page break

# Find the page break paragraph
for i in range(insert_idx + 1, len(tgt.paragraphs)):
    if tgt.paragraphs[i].text.strip() == "" and i > insert_idx:
        # Insert section header here
        hdr = tgt.paragraphs[i]
        # Turn it into a heading
        hdr.style = tgt.styles["Heading 1"]
        hr = hdr.add_run("ANNEXE A — Explication détaillée des cibles ML par hypothèse")
        hr.bold = True
        hr.font.size = Pt(15)
        hr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        break

# ── Save ──────────────────────────────────────────────────────────────────────

tgt.save(str(OUT))
print(f"Document fusionne sauvegarde : {OUT}")
print(f"Taille : {OUT.stat().st_size / 1024:.0f} Ko")
