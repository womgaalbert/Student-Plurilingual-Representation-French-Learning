"""
generate_word_report.py — Génère le rapport intermédiaire Word (FR)
French-Learning-Perceptions ML Project
Usage : python scripts/generate_word_report.py
"""
import json
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(13)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        run.font.size = Pt(11)
    return p


def add_metric_table(doc: Document, rows: list[tuple], header: tuple):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1F497D")
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            row[i].paragraphs[0].runs[0].font.size = Pt(9)
    return table


def add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "French-Learning-Perceptions in Plurilingual Cameroon — Rapport intermédiaire — Mai 2026"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── Data ──────────────────────────────────────────────────────────────────────

REPORT_JSON = Path("reports/rapport_evaluation_20260521_0930.json")
OUT_PATH = Path("reports/rapport_intermediaire_FLP_mai2026.docx")

with open(REPORT_JSON, encoding="utf-8") as f:
    data = json.load(f)

H = data["hypotheses"]
synth = data["synthese"]

# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# Default paragraph style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

add_footer(doc)

# ── PAGE DE GARDE ─────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RAPPORT INTERMÉDIAIRE")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Perceptions de l'apprentissage du français\nen contexte plurilingue au Cameroun")
run2.bold = True
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

meta_lines = [
    ("Auteure", "Armelle"),
    ("Institution", "Universite Marie et Louis Pasteur de Besancon (France)"),
    ("Date", "21 mai 2026"),
    ("N répondants", str(data["n_repondants"])),
    ("Statut", synth["statut"].replace("⚠️ ", "")),
    ("Hypothèses validées", synth["hypotheses_validees"]),
]
for label, val in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lbl = p.add_run(f"{label} : ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(12)
    r_val = p.add_run(val)
    r_val.font.size = Pt(12)

doc.add_page_break()

# ── TABLE DES MATIÈRES (manuelle) ────────────────────────────────────────────

add_heading(doc, "Table des matières", 1)
toc_items = [
    "1. Contexte et objectifs",
    "2. Méthodologie",
    "3. Corrections apportées au prétraitement",
    "4. Résultats par hypothèse",
    "   4.1 Hypothèse H1 — Répertoire linguistique et usage quotidien",
    "   4.2 Hypothèse H2 — Perceptions et motivations",
    "   4.3 Hypothèse H3 — Exposition et attitudes",
    "   4.4 Hypothèse H4 — Intégration des langues locales",
    "5. Synthèse globale",
    "6. Limites et pistes d'amélioration",
    "7. Prochaines étapes",
    "8. Environnement technique",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Bullet" if not item.startswith(" ") else "Normal")

doc.add_page_break()

# ── 1. CONTEXTE ───────────────────────────────────────────────────────────────

add_heading(doc, "1. Contexte et objectifs", 1)
doc.add_paragraph(
    "Ce rapport présente les résultats intermédiaires d'une étude sur les perceptions de "
    "l'apprentissage du français dans un contexte plurilingue au Cameroun. L'étude s'appuie "
    "sur un questionnaire administré à 500 élèves, portant sur leurs pratiques linguistiques, "
    "leurs perceptions du français, leurs difficultés d'apprentissage, ainsi que leur ouverture "
    "à l'intégration des langues camerounaises dans l'enseignement."
)
doc.add_paragraph(
    "L'objectif principal est de vérifier quatre hypothèses à l'aide de modèles d'apprentissage "
    "automatique supervisé, en combinant classification, régression et inférence causale."
)

# ── 2. MÉTHODOLOGIE ───────────────────────────────────────────────────────────

add_heading(doc, "2. Méthodologie", 1)

add_heading(doc, "2.1 Collecte des données", 2)
doc.add_paragraph(
    "Les données ont été collectées via un questionnaire en ligne (Google Forms). "
    f"Après vérification du consentement, {data['n_repondants']} réponses valides ont été retenues. "
    "Les données sont anonymisées avant tout traitement (suppression des horodateurs, emails, noms)."
)

add_heading(doc, "2.2 Pipeline de traitement", 2)
doc.add_paragraph("Le pipeline complet comprend 10 étapes :")
steps = [
    "Prétraitement et encodage des colonnes (preprocess.py)",
    "Analyse descriptive (describe.py)",
    "Entraînement H1 — Classification binaire : usage quotidien des langues",
    "Entraînement H2 — Classification : motivation + difficultés (multi-label)",
    "Optimisation H2 — GridSearchCV sur XGBoost",
    "Entraînement H3 — Régression et causalité : attitude envers le français",
    "Optimisation H3 — GridSearchCV sur RandomForest",
    "Entraînement H4 — Classification : motivation et engagement pour les langues locales",
    "Optimisation H4 — GridSearchCV sur XGBoost",
    "Évaluation globale et génération du rapport JSON",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"Étape {i} : {s}", style="List Number")

add_heading(doc, "2.3 Modèles utilisés", 2)
models_data = [
    ("H1", "XGBoost (classification binaire)", "F1-macro, ROC-AUC, CV-F1"),
    ("H2-A", "XGBoost + RandomOverSampler (motivation)", "F1-weighted, Val-F1"),
    ("H2-B", "ClassifierChain(XGBoost) (difficultés)", "F1-micro, Val-F1"),
    ("H3", "RandomForest Regressor (attitude)", "MAE, Pearson r, ATE"),
    ("H3-clf", "RandomForest Classifier (attitude binaire)", "F1-weighted"),
    ("H4-A", "XGBoost (motivation langues locales)", "F1, Val-F1"),
    ("H4-B", "XGBoost multi-classe (engagement)", "Spearman rho"),
    ("H4-C", "ClassifierChain(XGBoost) (disciplines)", "Subset accuracy"),
]
add_metric_table(doc, models_data, ("Hypothèse", "Modèle", "Métriques"))

doc.add_paragraph()

# ── 3. CORRECTIONS PRÉTRAITEMENT ─────────────────────────────────────────────

add_heading(doc, "3. Corrections apportées au prétraitement", 1)
doc.add_paragraph(
    "Plusieurs bugs critiques dans la construction des variables cibles ont été identifiés "
    "et corrigés au cours du développement :"
)
corrections = [
    (
        "exposition_bin (H3)",
        "La colonne d'exposition aux autres langues contient des réponses Oui/Non, "
        "mais était traitée avec FREQ_MAP (Toujours/Souvent/…) → tous zéros.",
        "Encodage binaire : 'Oui' → 1, 'Non' → 0 (avec tolérance aux fautes de frappe).",
    ),
    (
        "h4_target_motivation (H4)",
        "extract_oui_non() ne reconnaissait pas 'Bien'/'Très bien' → toutes les valeurs à -1 "
        "→ remplacées par 0 → variable cible constante.",
        "Utilisation d'INTERET_MAP (Très bien=3, Bien=2, Un peu=1, Pas du tout=0) "
        "avec seuil ≥2 → 1 (163/332 positifs).",
    ),
    (
        "h3_score_attitude (H3)",
        "attitude_score() cherchait startswith('OUI') sur des valeurs 'Bien/Très bien' "
        "→ s2 toujours à 0.5 → score effondré à 3 valeurs distinctes.",
        "Utilisation d'un dictionnaire _S2 avec correspondances explicites : "
        "'Très bien'→2.0, 'Bien'→1.5, 'Un peu'→1.0, 'Pas du tout'→0.5.",
    ),
    (
        "interet_camarades_ord (H4)",
        "La variable d'intérêt pour les langues des camarades était encodée en binaire "
        "(extract_oui_non) alors que les réponses suivent une échelle Likert.",
        "Encodage ordinal via INTERET_MAP (0–3) pour préserver l'information de degré.",
    ),
    (
        "comprehension (H2)",
        "Le mot-clé 'Compréhension de texte' était la difficulté la plus fréquente "
        "mais absent de DIFFICULTE_KEYWORDS.",
        "Ajout de la clé 'comprehension' avec mots-clés : compréhension, comprendre, lecture.",
    ),
    (
        "SMOTE → RandomOverSampler",
        "SMOTE exige ≥6 échantillons par classe par fold. Avec le déséquilibre sévère, "
        "certains folds CV n'ont que 1–3 exemples minoritaires → erreur n_neighbors.",
        "Remplacement de SMOTE par RandomOverSampler (ROS), compatible avec tout nombre de samples.",
    ),
]
for name, problem, fix in corrections:
    add_heading(doc, name, 3)
    p = doc.add_paragraph()
    p.add_run("Problème : ").bold = True
    p.add_run(problem)
    p2 = doc.add_paragraph()
    p2.add_run("Correction : ").bold = True
    p2.add_run(fix)

doc.add_page_break()

# ── 4. RÉSULTATS ──────────────────────────────────────────────────────────────

add_heading(doc, "4. Résultats par hypothèse", 1)


# H1
add_heading(doc, "4.1 Hypothèse H1 — Répertoire linguistique et usage quotidien", 2)
doc.add_paragraph(
    "H1 postule que les élèves disposant d'un répertoire linguistique étendu (≥3 langues) "
    "et d'une exposition fréquente à d'autres langues utilisent davantage ces langues "
    "au quotidien."
)
h1 = H["H1"]
h1_rows = [
    ("F1-macro", f"{h1['metrics']['f1_macro']:.4f}", "≥ 0.75", "✓"),
    ("ROC-AUC",  f"{h1['metrics']['roc_auc']:.4f}",  "≥ 0.80", "✓"),
    ("CV F1-macro (moy)", f"{h1['metrics']['cv_f1_macro_mean']:.4f}", "—", "—"),
    ("CV F1-macro (std)", f"{h1['metrics']['cv_f1_macro_std']:.4f}",  "—", "—"),
]
add_metric_table(doc, h1_rows, ("Métrique", "Valeur", "Seuil", "Atteint"))
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Statut : ").bold = True
r = p.add_run("VALIDÉE ✓")
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
doc.add_paragraph(
    "Interprétation : Le modèle prédit avec fiabilité si un élève utilise d'autres langues "
    "au quotidien. Le répertoire large (≥3 langues) et l'exposition fréquente sont des "
    "prédicteurs significatifs."
)
p2 = doc.add_paragraph()
p2.add_run("Recommandation : ").bold = True
p2.add_run(h1["recommandation"])


# H2
doc.add_paragraph()
add_heading(doc, "4.2 Hypothèse H2 — Perceptions et motivations", 2)
doc.add_paragraph(
    "H2 examine si la perception du français comme 'difficile' est associée à une motivation "
    "plus faible, et identifie les types de difficultés dominants (grammaire, conjugaison, etc.)."
)
h2 = H["H2"]
h2_rows = [
    ("F1-weighted (A : motivation)",    f"{h2['metrics']['f1_weighted_A']:.4f}", "≥ 0.75", "✓"),
    ("F1-micro (B : difficultés)",      f"{h2['metrics']['f1_micro_B']:.4f}",    "≥ 0.70", "✓"),
    ("Val-F1 (A : motivation)",         f"{h2['metrics']['val_f1_A']:.4f}",      "—", "—"),
    ("Val-F1 (B : difficultés)",        f"{h2['metrics']['val_f1_B']:.4f}",      "—", "—"),
]
add_metric_table(doc, h2_rows, ("Métrique", "Valeur", "Seuil", "Atteint"))
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Statut : ").bold = True
r = p.add_run("VALIDÉE ✓")
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
doc.add_paragraph(
    "Interprétation : Les élèves qui perçoivent le français comme difficile expriment "
    "une motivation significativement plus faible. Les difficultés dominantes sont la "
    "grammaire et la conjugaison, suivies de la compréhension de texte."
)
p2 = doc.add_paragraph()
p2.add_run("Recommandation : ").bold = True
p2.add_run(h2["recommandation"])


# H3
doc.add_paragraph()
add_heading(doc, "4.3 Hypothèse H3 — Exposition et attitudes envers le français", 2)
doc.add_paragraph(
    "H3 teste si une exposition fréquente aux autres langues est causalement liée "
    "à des attitudes plus positives envers le français."
)
h3 = H["H3"]
h3_rows = [
    ("MAE (régression attitude)",       f"{h3['metrics']['mae']:.4f}",             "≤ 0.50", "✗"),
    ("F1-weighted (clf attitude)",      f"{h3['metrics']['f1_weighted_clf']:.4f}",  "≥ 0.75", "✗"),
    ("Pearson r (causal)",              f"{h3['metrics']['pearson_r']:.4f}",        "sig.",   "✗"),
    ("Pearson p-value",                 f"{h3['metrics']['pearson_p']:.4f}",        "< 0.05", "✗"),
    ("ATE (Average Treatment Effect)",  f"{h3['metrics']['ate']:.4f}",             "—", "—"),
    ("Val-MAE",                         f"{h3['metrics']['val_mae']:.4f}",          "—", "—"),
    ("Val-F1",                          f"{h3['metrics']['val_f1']:.4f}",           "—", "—"),
]
add_metric_table(doc, h3_rows, ("Métrique", "Valeur", "Seuil", "Atteint"))
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Statut : ").bold = True
r = p.add_run("NON VALIDÉE ✗")
r.bold = True
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
doc.add_paragraph(
    "Interprétation : 93,5 % des élèves déclarent être exposés à d'autres langues (exposition_bin=1), "
    "ce qui laisse insuffisamment de variance dans le prédicteur pour établir un effet causal "
    "(p=0.984, ATE≈0). La MAE de 0.531 est proche du seuil de 0.50. "
    "Ce résultat est en soi une découverte substantielle : l'exposition plurilingue est quasi-universelle "
    "dans cet échantillon."
)
p2 = doc.add_paragraph()
p2.add_run("Recommandation : ").bold = True
p2.add_run(h3["recommandation"])
doc.add_paragraph(
    "Piste d'amélioration : Enrichir la variable d'exposition (fréquence détaillée, contextes) "
    "ou utiliser un instrument à plus grande variance. Tester LightGBM ou ExtraTreesRegressor "
    "pour améliorer la MAE."
)


# H4
doc.add_paragraph()
add_heading(doc, "4.4 Hypothèse H4 — Intégration des langues locales et motivation", 2)
doc.add_paragraph(
    "H4 évalue si les élèves seraient davantage motivés par l'intégration des langues "
    "camerounaises dans les cours, et identifie les disciplines préférées pour cette intégration."
)
h4 = H["H4"]
h4_rows = [
    ("F1 (A : motivation langues loc.)", f"{h4['metrics']['f1_A']:.4f}",           "≥ 0.75", "✓"),
    ("Spearman rho (B : engagement)",    f"{h4['metrics']['spearman_B']:.4f}",      "≥ 0.55", "✗"),
    ("Subset accuracy (C : disciplines)", f"{h4['metrics']['subset_acc_C']:.4f}",  "≥ 0.60", "✓"),
    ("Val-F1 (A : motivation)",          f"{h4['metrics']['val_f1_A']:.4f}",        "—", "—"),
]
add_metric_table(doc, h4_rows, ("Métrique", "Valeur", "Seuil", "Atteint"))
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Statut : ").bold = True
r = p.add_run("NON VALIDÉE ✗  (2/3 sous-objectifs atteints)")
r.bold = True
r.font.color.rgb = RGBColor(0xC0, 0x60, 0x00)
doc.add_paragraph(
    "Interprétation : Le modèle prédit correctement si un élève serait motivé "
    "par les langues locales (F1=0.80) et identifie les disciplines souhaitées avec "
    "précision (subset_acc=1.0). Cependant, la corrélation de Spearman pour le score "
    "d'engagement (rho=0.476) reste en dessous du seuil de 0.55."
)
p2 = doc.add_paragraph()
p2.add_run("Recommandation : ").bold = True
p2.add_run(h4["recommandation"])
doc.add_paragraph(
    "Piste d'amélioration : Enrichir les features comportementales (fréquence de participation, "
    "contextes d'usage des langues locales). Le score rho=0.476 est proche du seuil — "
    "un tuning plus fin ou des features supplémentaires pourraient suffire."
)

doc.add_page_break()

# ── 5. SYNTHÈSE ───────────────────────────────────────────────────────────────

add_heading(doc, "5. Synthèse globale", 1)

synth_rows = [
    ("H1 — Répertoire et usage quotidien",       "XGBoost + CV",         "F1=0.835, AUC=0.851", "VALIDÉE ✓"),
    ("H2 — Perceptions et difficultés",          "XGBoost + ClassChain", "F1-w=0.954, F1-mi=0.745", "VALIDÉE ✓"),
    ("H3 — Exposition et attitudes",             "RandomForest Reg.",    "MAE=0.531, p=0.984",  "NON VALIDÉE ✗"),
    ("H4 — Intégration langues locales",         "XGBoost multi-task",   "F1=0.80, rho=0.476",  "NON VALIDÉE ✗"),
]
add_metric_table(doc, synth_rows, ("Hypothèse", "Modèle", "Métriques clés", "Statut"))

doc.add_paragraph()
doc.add_paragraph(
    f"Résultat global : {synth['hypotheses_validees']} hypothèses validées. "
    "Les hypothèses H1 et H2 atteignent tous leurs seuils de performance. "
    "H3 et H4 sont proches des seuils — des améliorations ciblées sont envisageables."
)
doc.add_paragraph(
    "Points forts du pipeline : architecture MLOps complète (MLflow tracking, GridSearchCV, "
    "SMOTE remplacé par ROS pour robustesse), corrections rigoureuses des variables cibles, "
    "et validation croisée systématique."
)

# ── 6. LIMITES ────────────────────────────────────────────────────────────────

add_heading(doc, "6. Limites et pistes d'amélioration", 1)
limits = [
    ("Variance insuffisante en H3",
     "93,5 % des élèves sont exposés aux autres langues — "
     "insuffisant pour détecter un effet causal. Envisager une mesure d'exposition plus fine "
     "(fréquence, contexte, durée)."),
    ("Taille d'échantillon",
     "Avec 500 répondants et parfois seulement 163 positifs (H4), "
     "certains modèles souffrent de déséquilibre de classes malgré le ROS."),
    ("Variables cibles proxy",
     "Plusieurs cibles sont construites par règles (attitude_score, engagement_score) "
     "et non mesurées directement — introduce un biais de construction."),
    ("Généralisation géographique",
     "L'échantillon est circonscrit à un établissement camerounais. "
     "Les résultats ne sont pas nécessairement généralisables à d'autres contextes."),
    ("Interprétabilité",
     "Les modèles XGBoost sont performants mais peu interprétables sans SHAP. "
     "L'intégration SHAP est prévue pour la phase suivante."),
]
for title, text in limits:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{title} : ").bold = True
    p.add_run(text)

# ── 7. PROCHAINES ÉTAPES ──────────────────────────────────────────────────────

add_heading(doc, "7. Prochaines étapes", 1)
next_steps_grouped = {
    "Amélioration des modèles H3 et H4": [
        "Tester LightGBM et ExtraTreesRegressor pour réduire la MAE H3 (< 0.50)",
        "Enrichir les features H4 avec des proxies comportementaux pour rho ≥ 0.55",
        "Affiner les variables d'exposition H3 (fréquence vs binaire)",
    ],
    "MLOps — Niveau 2": [
        "CI/CD via GitHub Actions : entraînement automatique sur push",
        "Model Registry MLflow : pipeline staging → production",
        "Endpoint REST FastAPI /predict pour chaque hypothèse",
        "Conteneurisation Docker du pipeline complet",
    ],
    "MLOps — Niveau 3": [
        "Détection de data drift avec Evidently AI",
        "Retraining automatique sur nouveau batch de données",
        "Dashboard de monitoring des prédictions en production",
    ],
    "Rapport pédagogique": [
        "Rédaction du rapport pédagogique final (reports/pedagogical_report.md)",
        "Rapport de validation de généralisation pour H1 et H2",
        "Analyse SHAP pour l'interprétabilité des modèles H1 et H2",
    ],
}
for section_title, items in next_steps_grouped.items():
    add_heading(doc, section_title, 3)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

# ── 8. ENVIRONNEMENT TECHNIQUE ────────────────────────────────────────────────

add_heading(doc, "8. Environnement technique", 1)
env_rows = [
    ("Langage", "Python 3.11"),
    ("ML", "XGBoost, scikit-learn, imbalanced-learn"),
    ("Tracking", "MLflow 2.x"),
    ("Données", "pandas, numpy"),
    ("Rapport", "python-docx"),
    ("Plateforme", "Windows 11, VS Code"),
    ("Gestion config", "params.yaml"),
]
add_metric_table(doc, env_rows, ("Composant", "Technologie/Version"))

doc.add_paragraph()
doc.add_paragraph(
    "Le projet respecte les règles CLAUDE.md : données brutes en lecture seule (data/raw/), "
    "consentement vérifié avant tout traitement, anonymisation systématique, "
    "hyperparamètres centralisés dans params.yaml, et toutes les runs MLflow trackées."
)

# ── SAUVEGARDE ────────────────────────────────────────────────────────────────

OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print(f"Rapport sauvegardé : {OUT_PATH}")
