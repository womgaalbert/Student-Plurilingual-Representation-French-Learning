"""
generate_mlops_audit.py — MLOps Audit Report (Word FR)
French-Learning-Perceptions ML Project
Usage : python scripts/generate_mlops_audit.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path("reports/audit_mlops_mai2026.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    colors = {1: "1F497D", 2: "2E74B5", 3: "404040"}
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(colors.get(level, "000000"))
    run.font.size = Pt({1: 16, 2: 13, 3: 11}.get(level, 11))
    return p


def table(doc, rows, header, col_widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        cell_bg(hdr[i], "1F497D")
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        if col_widths:
            hdr[i].width = Inches(col_widths[i])
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            p = row[i].paragraphs[0]
            if p.runs:
                p.runs[0].font.size = Pt(9)
            if col_widths:
                row[i].width = Inches(col_widths[i])
    return t


def badge(doc, score_pct, label):
    p = doc.add_paragraph()
    if score_pct >= 80:
        icon, color = "🟢", "007000"
    elif score_pct >= 40:
        icon, color = "🟡", "C06000"
    else:
        icon, color = "🔴", "C00000"
    r = p.add_run(f"{icon} {label} — {score_pct}% ")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string(color)
    return p


def footer(doc):
    sec = doc.sections[0]
    ft = sec.footer
    p = ft.paragraphs[0]
    p.text = "Audit MLOps — French-Learning-Perceptions in Plurilingual Cameroon — Mai 2026"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(2.5)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)
footer(doc)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AUDIT MLOps")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("French-Learning-Perceptions in Plurilingual Cameroon")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Évaluation de la plateforme ML selon 10 dimensions\n"
              "d'un système MLOps de qualité production")
r.font.size = Pt(12)

doc.add_paragraph()
for label, val in [
    ("Date de l'audit", "10 juin 2026"),
    ("Auditeur", "Albert Womga — AI/ML Engineering"),
    ("Chercheuse", "Chancelline Armelle Nongni Kendjio"),
    ("Périmètre", "Pipeline ML complet : données → modèles → évaluation"),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{label} : ").bold = True
    p.add_run(val)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SYNTHÈSE EXÉCUTIVE
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Synthèse exécutive", 1)

doc.add_paragraph(
    f"Cet audit évalue la maturité MLOps du projet French-Learning-Perceptions selon "
    f"10 dimensions critiques pour un passage en production. L'évaluation couvre "
    f"l'état actuel (mai 2026), identifie les forces et les lacunes, et fournit "
    f"des recommandations priorisées (P0 = critique, P1 = important, P2 = souhaitable)."
)

doc.add_paragraph()
table(doc, [
    ("📁 Données",              "75%", "🟡", "Schéma défini, contrats absents"),
    ("⚙️ Pipelines",             "70%", "🟡", "Orchestré mais monolithique"),
    ("🗂️ Modèles",               "45%", "🟡", "Versionnés fichiers, pas de Registry"),
    ("🚀 Serving",              "0%",  "🔴", "Aucune infrastructure de serving"),
    ("📊 Monitoring",            "0%",  "🔴", "Aucun monitoring de production"),
    ("🧪 Expérimentations",      "85%", "🟢", "MLflow complet, 116 runs, 11 exp."),
    ("🛡️ Gouvernance",           "35%", "🔴", "Consentement OK, pas de fairness"),
    ("🔄 CI/CD",                 "40%", "🟡", "CI tests, pas de train auto"),
    ("✅ Tests",                 "55%", "🟡", "39 passent, 3 cassés, pas d'intégration"),
    ("📚 Documentation",         "65%", "🟡", "CLAUDE.md excellent, pas de runbooks"),
], ("Dimension", "Score", "Niveau", "Résumé"))

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Score global MLOps : 47% ").bold = True
p.add_run("— Niveau 1 consolidé, chemin vers les Niveaux 2 et 3 identifié.")

heading(doc, "Top 5 actions critiques (P0)", 2)
for action in [
    "1. Servir les modèles via FastAPI + Docker (dimension 🚀 — actuellement 0%)",
    "2. Activer le Model Registry MLflow avec staging/production (dimension 🗂️)",
    "3. Implémenter la détection de drift via Evidently AI (dimension 📊)",
    "4. Ajouter le pipeline d'entraînement automatique en CI/CD (dimension 🔄)",
    "5. Corriger les 3 tests cassés + ajouter des tests d'intégration (dimension ✅)",
]:
    doc.add_paragraph(action, style="List Bullet")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# AUDIT DÉTAILLÉ PAR DIMENSION
# ══════════════════════════════════════════════════════════════════════════════

# ── 1. DATA ───────────────────────────────────────────────────────────────────

heading(doc, "1. 📁 Données — Score : 75% 🟡", 1)

table(doc, [
    ("data/raw/ en lecture seule",              "✅", "Règle absolue CLAUDE.md"),
    ("data/processed/ avec .gitkeep",           "✅", "6 CSV (clean + H1-H4)"),
    ("Filtrage consentement",                   "✅", "preprocess.py → 495/500"),
    ("Anonymisation systématique",              "✅", "Horodateur + identifiants supprimés"),
    ("Schéma de colonnes documenté",            "✅", "constants.py + CSV_COLUMN_MAP"),
    ("Contrats de données (schema validation)",  "❌", "Pas de Great Expectations / Pandera"),
    ("Versionnement des données (DVC)",          "❌", "Pas de data versioning"),
    ("Détection de schema drift",                "❌", "Pas d'alerte si nouvelle colonne"),
    ("Séparation train/val/test stricte",        "✅", "70/15/15 stratifié, seed=42"),
    ("Données de test hermétiques",              "✅", "X_test jamais vu en entraînement"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "Forces : La chaîne de traitement des données est robuste. Le consentement est "
    "vérifié, l'anonymisation est automatique, et la séparation train/val/test est "
    "stricte avec stratification. Les noms de colonnes sont centralisés dans constants.py."
)
doc.add_paragraph(
    "Lacunes : Absence de validation de schéma automatisée. Si le CSV source change "
    "(colonnes renommées, nouveau format), rien ne le détectera avant l'erreur runtime. "
    "Pas de data versioning (DVC) — impossible de reproduire exactement un run avec "
    "les données d'il y a 3 mois si le raw est mis à jour."
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P1] Ajouter Pandera ou Great Expectations pour valider le schéma d'entrée",
    "[P2] Versionner data/processed/ avec DVC pour la reproductibilité",
    "[P2] Ajouter un test de non-regression sur le nombre de colonnes",
]:
    doc.add_paragraph(a, style="List Bullet")


# ── 2. PIPELINES ──────────────────────────────────────────────────────────────

heading(doc, "2. ⚙️ Pipelines — Score : 70% 🟡", 1)

table(doc, [
    ("Orchestration 10 étapes",          "✅", "pipeline.py : preprocess → H1→H4 → evaluate"),
    ("Étapes modulaires et indépendantes","✅", "Chaque train_h*() est appelable seule"),
    ("Configuration externalisée",       "✅", "params.yaml : zéro hardcoding"),
    ("Reproductibilité (seed fixe)",     "✅", "random_state=42 partout"),
    ("Rerun partiel possible",           "⚠️", "Via --hypothesis mais pas de skip si déjà fait"),
    ("Gestion des dépendances inter-étapes","⚠️","Pas de DAG — exécution séquentielle imposée"),
    ("Parallélisation des étapes",       "❌", "H1-H4 entraînés séquentiellement"),
    ("Orchestrateur externe (Airflow)",  "❌", "Pas d'intégration Prefect/Airflow/Dagster"),
    ("Gestion des échecs + retry",       "❌", "Si une étape crash, tout le pipeline s'arrête"),
    ("Cache des étapes intermédiaires",  "❌", "CamemBERT rechargé 2× (H3 puis H4)"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "Forces : Le pipeline est bien structuré et modulaire. Chaque entraînement H1-H4 "
    "peut être appelé indépendamment. La configuration params.yaml garantit la "
    "reproductibilité. Les 10 étapes sont clairement documentées."
)
doc.add_paragraph(
    "Lacunes : Les modèles sont entraînés séquentiellement alors que H1, H2, H3, H4 "
    "pourraient être parallélisés (gain ~3-4×). Aucune gestion de cache — CamemBERT "
    "est rechargé 2 fois. Pas de mécanisme de retry en cas d'échec."
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P1] Paralléliser H1/H2/H3/H4 (joblib ou ProcessPoolExecutor)",
    "[P1] Partager le modèle CamemBERT entre H3 et H4 (chargement unique)",
    "[P2] Ajouter un flag --skip-existing pour éviter de ré-entraîner",
    "[P2] Migrer vers Prefect ou Dagster pour le DAG management",
]:
    doc.add_paragraph(a, style="List Bullet")


# ── 3. MODELS ─────────────────────────────────────────────────────────────────

heading(doc, "3. 🗂️ Modèles — Score : 45% 🟡", 1)

table(doc, [
    ("Sérialisation .pkl par run",       "✅", "models/h{n}/h{n}_{model}_{ts}.pkl"),
    ("MLflow logging (params, metrics)",  "✅", "Chaque run loggue params + métriques"),
    ("MLflow artefacts (.pkl uploadé)",   "✅", "log_model_artifact() systématique"),
    ("Model Registry (staging→prod)",     "❌", "0 modèles enregistrés dans le Registry"),
    ("Traçabilité (data → code → modèle)","⚠️", "MLflow trace le code, pas les données"),
    ("Version sémantique des modèles",    "❌", "Horodatage uniquement, pas de semver"),
    ("Baseline model systématique",       "⚠️", "DummyClassifier présent mais peu utilisé"),
    ("Stockage centralisé (S3/GCS)",      "❌", "Stockage local uniquement"),
    ("Rollback facilité",                 "❌", "Pas de mécanisme de rollback"),
    ("Nettoyage des modèles obsolètes",   "❌", "147 .pkl accumulés sans purge"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "Forces : 147 modèles sérialisés avec horodatage. MLflow trace tous les paramètres, "
    "métriques et artefacts. Les tags 'status' (✅/⚠️) facilitent le filtrage dans l'UI."
)
doc.add_paragraph(
    "Lacunes : Le Model Registry MLflow n'est pas activé — 0 modèles enregistrés. "
    "Sans Registry, impossible de savoir quel modèle est en production, quel dataset "
    "l'a entraîné, ou comment revenir à une version antérieure. Les 147 fichiers .pkl "
    "s'accumulent sans politique de rétention."
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P0] Activer le Model Registry MLflow → staging/production/archived",
    "[P0] Enregistrer le meilleur modèle de chaque hypothèse avec métadonnées",
    "[P1] Ajouter le lineage données (hash du CSV d'entraînement dans les tags MLflow)",
    "[P1] Implémenter une politique de rétention (garder 3 derniers par hypothèse)",
    "[P2] Versionner avec semver : v1.0.0-h1, v1.0.0-h2, etc.",
]:
    doc.add_paragraph(a, style="List Bullet")


# ── 4. SERVING ────────────────────────────────────────────────────────────────

heading(doc, "4. 🚀 Serving — Score : 0% 🔴", 1)

table(doc, [
    ("API REST (FastAPI/Flask)",     "❌", "api/ vide — aucun endpoint"),
    ("Batch inference",              "❌", "Pas de pipeline batch"),
    ("Streaming / event-driven",     "❌", "Pas de support streaming"),
    ("Dockerfile",                   "❌", "Pas de conteneurisation"),
    ("docker-compose",               "❌", "Pas d'orchestration multi-services"),
    ("Health endpoint",              "❌", "Pas de /health"),
    ("Load testing",                 "❌", "Pas de test de charge"),
    ("GPU support",                  "N/A","CamemBERT sur CPU — pas nécessaire"),
    ("Authentication",               "❌", "Pas d'auth sur les endpoints"),
    ("Rate limiting",                "❌", "Pas de limitation de débit"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "C'est la dimension la plus critique — 0%. Aucun modèle n'est serviable en l'état. "
    "Pour passer du stade 'recherche' au stade 'production', il faut exposer chaque "
    "hypothèse via une API REST. FastAPI est le choix naturel (déjà dans les dépendances)."
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P0] Créer api/main.py avec 4 endpoints : POST /predict/h1, /h2, /h3, /h4",
    "[P0] Écrire le Dockerfile (python:3.11-slim + modèles .pkl)",
    "[P0] Ajouter docker-compose.yml (API + MLflow server)",
    "[P1] Ajouter /health et /metrics endpoints",
    "[P1] Documenter le schéma d'entrée/sortie de chaque endpoint",
    "[P2] Ajouter authentification par token",
    "[P2] Load testing avec locust",
]:
    doc.add_paragraph(a, style="List Bullet")


# ── 5. MONITORING ─────────────────────────────────────────────────────────────

heading(doc, "5. 📊 Monitoring — Score : 0% 🔴", 1)

table(doc, [
    ("Data drift detection",         "❌", "Pas d'Evidently AI ni de monitoring"),
    ("Model performance monitoring", "❌", "Pas de suivi des métriques en prod"),
    ("Alerting (Slack/email)",       "❌", "Pas de système d'alerte"),
    ("Prediction logging",           "❌", "Pas de log des prédictions"),
    ("Dashboard temps réel",         "❌", "Pas de Grafana"),
    ("Concept drift detection",      "❌", "Pas de détection de drift conceptuel"),
    ("Feature distribution tracking","❌", "Pas de suivi de distribution"),
    ("Outlier detection en production","❌","Pas de détection d'anomalies"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "Deuxième dimension critique à 0%. Les modèles se dégradent silencieusement en "
    "production. Sans monitoring, un modèle peut produire des prédictions erronées "
    "pendant des semaines sans que personne ne le sache. Pour un projet de recherche "
    "avec valeur sociale (recommandations pédagogiques), c'est inacceptable en production."
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P0] Intégrer Evidently AI pour le data drift sur les features d'entrée",
    "[P0] Logger chaque prédiction avec timestamp + features + résultat",
    "[P1] Créer monitoring/drift_detector.py avec seuils d'alerte configurables",
    "[P1] Dashboard Grafana : distribution des features, volume de requêtes, latence",
    "[P1] Alerte Slack/email si drift > seuil sur 2 fenêtres consécutives",
    "[P2] Métriques de fairness en production (démographie des prédictions)",
]:
    doc.add_paragraph(a, style="List Bullet")


# ── 6. EXPERIMENTS ────────────────────────────────────────────────────────────

heading(doc, "6. 🧪 Expérimentations — Score : 85% 🟢", 1)

table(doc, [
    ("MLflow tracking",               "✅", "11 expériences, 116 runs"),
    ("Nested runs (parent + enfants)","✅", "Pipeline parent → H1-H4 enfants"),
    ("Paramètres loggés",             "✅", "log_params_from_config() systématique"),
    ("Métriques loggées",             "✅", "Toutes les métriques par hypothèse"),
    ("Artefacts sauvegardés",         "✅", "Modèles .pkl + rapports JSON"),
    ("Tags de statut (✅/⚠️)",        "✅", "mlflow.set_tag('status', ...)"),
    ("Comparaison visuelle (UI)",     "✅", "MLflow UI sur port 5000"),
    ("GridSearchCV loggé",            "✅", "best_params_ enregistrés dans MLflow"),
    ("Reproductibilité complète",     "⚠️", "seed=42 mais pas de hash des données"),
    ("A/B testing framework",         "❌", "Pas de cadre de test A/B"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "C'est la dimension la plus mature du projet. Avec 11 expériences distinctes, "
    "116 runs trackés et des nested runs pour le pipeline, le tracking MLflow est "
    "exemplaire pour un projet de cette taille. Chaque run enregistre ses paramètres, "
    "métriques, artefacts et tags de statut."
)
doc.add_paragraph(
    "Lacune mineure : pas de hash des données d'entraînement, ce qui empêche "
    "la reproductibilité parfaite si les données sous-jacentes changent."
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P1] Ajouter le hash MD5 du CSV d'entraînement dans les tags MLflow",
    "[P1] Migrer vers un backend base de données (sqlite:///mlflow.db) — le backend "
         "fichier est déprécié depuis février 2026",
    "[P2] Ajouter un notebook de comparaison automatique des runs",
]:
    doc.add_paragraph(a, style="List Bullet")


# ── 7. GOVERNANCE ─────────────────────────────────────────────────────────────

heading(doc, "7. 🛡️ Gouvernance — Score : 35% 🔴", 1)

table(doc, [
    ("Consentement vérifié",          "✅", "Filtrage 'J'accepte' obligatoire"),
    ("Anonymisation",                 "✅", "Suppression identifiants avant export"),
    ("CODEOWNERS",                    "✅", ".github/CODEOWNERS présent"),
    ("Data read-only (raw/)",         "✅", "Règle absolue respectée"),
    ("Model cards",                   "❌", "Pas de documentation par modèle"),
    ("Fairness / bias evaluation",    "❌", "Pas d'analyse de biais démographique"),
    ("Audit trail",                   "⚠️", "MLflow = audit partiel, pas d'accès RBAC"),
    ("GDPR / data privacy",           "⚠️", "Anonymisation OK, pas de politique retention"),
    ("Explainability (SHAP)",         "⚠️", "Importé mais non utilisé dans les rapports"),
    ("Access control (RBAC)",         "❌", "Pas de contrôle d'accès"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "Forces : Le consentement et l'anonymisation sont solides, essentiels pour un "
    "projet traitant des données d'élèves mineurs. Le fichier CODEOWNERS est en place."
)
doc.add_paragraph(
    "Lacunes : Pas de model cards documentant les limites, biais et usage prévu de "
    "chaque modèle. SHAP est importé mais jamais utilisé en production de rapports. "
    "Pas d'analyse de fairness — le modèle H1 ne favorise-t-il pas certaines régions ?"
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P1] Rédiger une model card par hypothèse (limites, biais, usage prévu)",
    "[P1] Analyse SHAP sur H1 et H2 pour identifier les features dominantes",
    "[P1] Test de fairness : F1-score par région et par genre",
    "[P2] Définir une politique de rétention des données",
    "[P2] Ajouter un registre des décisions (ADR) dans docs/adr/",
]:
    doc.add_paragraph(a, style="List Bullet")


# ── 8. CI/CD ──────────────────────────────────────────────────────────────────

heading(doc, "8. 🔄 CI/CD — Score : 40% 🟡", 1)

table(doc, [
    ("GitHub Actions CI",            "✅", "ci.yml + tests.yml sur push/PR"),
    ("Validation params.yaml",       "✅", "Vérification automatisée dans CI"),
    ("Tests automatisés",            "✅", "pytest lancé dans CI"),
    ("Train automatique sur push",   "❌", "Pas de pipeline d'entraînement auto"),
    ("Déploiement automatique",      "❌", "Pas de déploiement continu"),
    ("Gates de déploiement",         "❌", "Pas de seuils bloquants avant déploiement"),
    ("Matrix testing (OS/Python)",   "⚠️", "windows-latest + ubuntu-latest"),
    ("Secrets management",           "❌", "Pas de GitHub Secrets pour tokens"),
    ("Multi-branches (dev/staging)", "❌", "CI sur main/dev seulement"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "Forces : Deux workflows GitHub Actions fonctionnels. La validation de params.yaml "
    "dans le CI est une bonne pratique. Le cache pip est activé."
)
doc.add_paragraph(
    "Lacunes : Pas de pipeline d'entraînement automatique — le modèle n'est ré-entraîné "
    "que manuellement. Pas de déploiement continu ni de gates bloquantes (ex: 'ne pas "
    "déployer si F1 < 0.70'). Pas de gestion des secrets pour les tokens HF."
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P1] Ajouter .github/workflows/train.yml : entraînement auto sur push main",
    "[P1] Ajouter des gates : bloquer le déploiement si métriques < seuils",
    "[P1] Configurer GitHub Secrets pour HF_TOKEN (téléchargement CamemBERT)",
    "[P2] Ajouter un workflow de déploiement (build Docker + push registry)",
]:
    doc.add_paragraph(a, style="List Bullet")


# ── 9. TESTS ──────────────────────────────────────────────────────────────────

heading(doc, "9. ✅ Tests — Score : 55% 🟡", 1)

table(doc, [
    ("Tests unitaires",              "✅", "39 passés sur 42"),
    ("Tests de preprocessing",       "✅", "Consentement, anonymisation, age"),
    ("Tests des constantes",         "✅", "FREQ_MAP, LIKERT_MAP coverage"),
    ("Tests de range des cibles",    "✅", "h3_score ∈ [1,5], h4_engagement ∈ {1,2,3,4}"),
    ("Tests de non-leakage",         "✅", "Targets absentes des features"),
    ("Tests d'intégration",          "❌", "Pas de test end-to-end du pipeline"),
    ("Tests de régression modèle",   "❌", "Pas de test 'le modèle ne doit pas régresser'"),
    ("Couverture de code",           "⚠️", "Pas mesurée systématiquement"),
    ("Tests de performance",         "❌", "Pas de benchmark de latence"),
    ("3 tests cassés",               "⚠️", "Non mis à jour après renommage colonnes"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "Forces : 39 tests passent, couvrant le preprocessing, les constantes, les ranges "
    "des cibles et l'absence de data leakage. La structure pytest est propre."
)
doc.add_paragraph(
    "Lacunes : 3 tests échouent car ils référencent d'anciens noms de colonnes "
    "(exposition_freq → exposition_bin, perc_difficile → perc_difficil tronqué à 8 chars). "
    "Aucun test d'intégration ne valide le pipeline complet. Pas de tests de non-régression "
    "pour les modèles."
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P0] Corriger les 3 tests cassés (mise à jour des noms de colonnes)",
    "[P1] Ajouter un test d'intégration : pipeline complet sur un subset de 10 lignes",
    "[P1] Ajouter un test de non-régression : F1 H1 ≥ 0.70, AUC ≥ 0.75",
    "[P1] Mesurer et suivre la couverture de code (pytest-cov ≥ 60%)",
    "[P2] Ajouter un test de latence : predict() < 500ms",
]:
    doc.add_paragraph(a, style="List Bullet")


# ── 10. DOCS ──────────────────────────────────────────────────────────────────

heading(doc, "10. 📚 Documentation — Score : 65% 🟡", 1)

table(doc, [
    ("CLAUDE.md complet",            "✅", "12 phases, 4 hypothèses, bugs, rules"),
    ("README.md",                    "✅", "Présentation projet + setup"),
    ("CONTRIBUTING.md",              "✅", "Guide de contribution"),
    ("SECURITY.md",                  "✅", "Politique de sécurité"),
    ("Docstrings en français",       "✅", "Toutes les fonctions publiques"),
    ("params.yaml documenté",        "✅", "Commentaires par section"),
    ("Rapports Word automatisés",    "✅", "2 scripts de génération"),
    ("Runbooks / procédures",        "❌", "Pas de guide 'que faire si X plante'"),
    ("Architecture Decision Records", "❌", "Pas d'ADRs"),
    ("Documentation API",            "❌", "Pas encore (API non créée)"),
], ("Critère", "Statut", "Détail"))

doc.add_paragraph()
doc.add_paragraph(
    "Forces : CLAUDE.md est un document de référence exceptionnel — 300+ lignes couvrant "
    "le pipeline complet, les checklists, les bugs corrigés et les règles absolues. "
    "Les docstrings sont systématiques. Les rapports Word sont générés automatiquement."
)
doc.add_paragraph(
    "Lacunes : Pas de runbooks opérationnels. Si le pipeline échoue à l'étape 7, "
    "que doit faire l'ingénieur qui reprend le projet ? Pas d'ADRs documentant "
    "les décisions architecturales (pourquoi ClassifierChain plutôt que MultiOutput ?)."
)

p = doc.add_paragraph()
p.add_run("Actions recommandées :").bold = True
for a in [
    "[P1] Écrire docs/runbook.md : procédures de debugging et reprise",
    "[P1] Créer docs/adr/ avec 4 ADRs : choix ClassifierChain, SMOTE→ROS, CamemBERT PCA, ExtraTrees vs RF",
    "[P2] Générer la documentation API automatiquement (FastAPI → /docs Swagger)",
    "[P2] Ajouter un diagramme d'architecture (Mermaid ou draw.io)",
]:
    doc.add_paragraph(a, style="List Bullet")


doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PLAN D'ACTION PRIORISÉ
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Plan d'action priorisé", 1)

heading(doc, "🔥 P0 — Bloquant pour la production (Sprint 1-2)", 2)
table(doc, [
    ("P0", "🚀", "Créer api/main.py avec endpoints /predict/h1-h4", "0% → 60%"),
    ("P0", "🚀", "Dockerfile + docker-compose.yml", "Infrastructure"),
    ("P0", "🗂️", "Activer Model Registry MLflow (staging/prod)", "45% → 70%"),
    ("P0", "📊", "Implémenter Evidently AI drift detection", "0% → 40%"),
    ("P0", "✅", "Corriger les 3 tests cassés", "55% → 65%"),
], ("Prio", "Dim.", "Action", "Impact"))

heading(doc, "🟡 P1 — Important (Sprint 3-4)", 2)
table(doc, [
    ("P1", "🔄", "Workflow GitHub Actions train.yml (auto-train)", "40% → 60%"),
    ("P1", "⚙️", "Paralléliser H1-H4 + cache CamemBERT", "70% → 85%"),
    ("P1", "🗂️", "Migrer MLflow backend → sqlite:///mlflow.db", "Maintenance"),
    ("P1", "📊", "Logger prédictions + dashboard Grafana", "0% → 40%"),
    ("P1", "🛡️", "Model cards + analyse SHAP + fairness test", "35% → 55%"),
    ("P1", "📚", "Rédiger docs/runbook.md + ADRs", "65% → 80%"),
    ("P1", "✅", "Tests d'intégration + non-régression modèle", "55% → 70%"),
    ("P1", "📁", "Validation de schéma avec Pandera", "75% → 85%"),
], ("Prio", "Dim.", "Action", "Impact"))

heading(doc, "🟢 P2 — Souhaitable (Backlog)", 2)
table(doc, [
    ("P2", "📁", "Versionnement données avec DVC", "75% → 90%"),
    ("P2", "⚙️", "Migrer vers Prefect/Dagster", "70% → 85%"),
    ("P2", "🗂️", "Semver des modèles + politique rétention", "45% → 60%"),
    ("P2", "🚀", "Auth endpoints + rate limiting + load testing", "60% → 80%"),
    ("P2", "📊", "Métriques fairness en production", "40% → 60%"),
    ("P2", "🛡️", "Politique rétention données + RBAC", "35% → 55%"),
    ("P2", "📚", "Diagramme architecture + doc API Swagger", "65% → 85%"),
    ("P2", "✅", "Benchmark de latence + stress tests", "55% → 70%"),
], ("Prio", "Dim.", "Action", "Impact"))


doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCORE GLOBAL
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Score global et projection", 1)

table(doc, [
    ("📁 Données",              "75%", "85%", "90%"),
    ("⚙️ Pipelines",             "70%", "85%", "90%"),
    ("🗂️ Modèles",               "45%", "70%", "85%"),
    ("🚀 Serving",              "0%",  "60%", "85%"),
    ("📊 Monitoring",            "0%",  "45%", "75%"),
    ("🧪 Expérimentations",      "85%", "90%", "95%"),
    ("🛡️ Gouvernance",           "35%", "55%", "70%"),
    ("🔄 CI/CD",                 "40%", "65%", "80%"),
    ("✅ Tests",                 "55%", "70%", "80%"),
    ("📚 Documentation",         "65%", "80%", "90%"),
    ("🏆 GLOBAL",               "47%", "70%", "84%"),
], ("Dimension", "Actuel", "Après P0+P1", "Cible"))

doc.add_paragraph()
doc.add_paragraph(
    "Le projet est actuellement à 47% de maturité MLOps — un score honorable pour "
    "un projet en phase de recherche. L'exécution des actions P0 et P1 porterait "
    "ce score à 70%, suffisant pour un déploiement en production supervisé. "
    "La cible de 84% correspond à un système MLOps de niveau 3, entièrement "
    "automatisé et monitoré."
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Verdict : ").bold = True
p.add_run(
    "Le pipeline de recherche est mature et bien structuré. "
    "Le gap principal se situe dans le serving (0%), le monitoring (0%) et le "
    "Model Registry (0 modèles). Ces trois dimensions sont le prérequis minimal "
    "pour tout déploiement en production. Une fois comblées, le projet pourra "
    "passer de 'excellent projet de recherche' à 'système ML opérationnel'."
)

# ── Footer ────────────────────────────────────────────────────────────────────

doc.save(str(OUT))
print(f"Audit sauvegardé : {OUT} ({OUT.stat().st_size / 1024:.0f} Ko)")
